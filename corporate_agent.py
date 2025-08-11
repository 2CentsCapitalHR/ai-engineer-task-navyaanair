"""
Corporate Agent — ADGM-compliant document reviewer (full implementation)



Outputs:
    - Downloadable reviewed DOCX per submitted file (with inline bracketed comments and summary page)
    - JSON structured report of findings and missing checklist items
"""

import os
import json
import tempfile
import shutil
from pathlib import Path
from typing import List, Dict, Any, Tuple

import gradio as gr
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt
from tqdm.auto import tqdm

# Chroma / Ollama imports
import chromadb
from chromadb.config import Settings
from langchain_ollama import OllamaEmbeddings, OllamaLLM
  # uses langchain-ollama package



# ---------------- CONFIG ----------------
CHROMA_PERSIST_DIR = "./chroma_langchain"   # update if your chroma persistence folder differs
EMBED_MODEL = "nomic-embed-text"            # ollama embedding model you pulled
LLM_MODEL = os.environ.get("LLM_MODEL", "gemma3:1b")  # adjust to the text-generation model you pulled
RAG_TOP_K = 4

# Checklist mapping for Company Incorporation (example)
CHECKLISTS = {
    "Company Incorporation": [
        "Articles of Association",
        "Memorandum of Association",
        "Board Resolution",
        "Incorporation Application Form",
        "Register of Members and Directors"
    ],
    # add other processes as needed
}

# Keywords to identify doc types (simple heuristics)
DOC_KEYWORDS = {
    "Articles of Association": ["articles of association", "articles of association (aoa)", "article of association"],
    "Memorandum of Association": ["memorandum of association", "memorandum of understanding", "moa", "mou"],
    "Board Resolution": ["board resolution", "resolution of the board"],
    "Incorporation Application Form": ["incorporation application", "application for incorporation"],
    "Register of Members and Directors": ["register of members", "register of directors", "register of members and directors"]
}

# ---------------- HELPERS: Chroma & RAG ----------------


def init_chroma_client(persist_dir: str = CHROMA_PERSIST_DIR):
    if not os.path.exists(persist_dir):
        print(f"Creating Chroma persist directory: {persist_dir}")
        os.makedirs(persist_dir)

    client = chromadb.PersistentClient(path=persist_dir)

    # Create collection if missing
    try:
        collection = client.get_collection("AllReferences")
    except chromadb.errors.NotFoundError:
        collection = client.create_collection("AllReferences")

    return client



def rag_retrieve(query: str, top_k: int = RAG_TOP_K) -> List[Dict[str, Any]]:
    """
    Retrieve top_k text chunks from the persisted Chroma index using Ollama embeddings.
    Returns list of dicts with keys: id, document (text), metadata, distance
    """
    client = init_chroma_client()
    collection = client.get_collection("AllReferences")
    embedder = OllamaEmbeddings(model=EMBED_MODEL)
    qvec = embedder.embed_query(query)
    results = collection.query(query_embeddings=[qvec], n_results=top_k, include=["metadatas", "documents", "distances"])
    outs = []
    if results and "ids" in results:
        for i in range(len(results["ids"][0])):
            outs.append({
                "id": results["ids"][0][i],
                "document": results["documents"][0][i],
                "metadata": results["metadatas"][0][i],
                "distance": results["distances"][0][i]
            })
    return outs

# ---------------- HELPERS: Text extraction & docx comment insertion ----------------
def extract_text_from_docx_bytes(file_bytes: bytes) -> str:
    """
    Load a .docx file from bytes and return the concatenated text.
    """
    # save temporarily and use python-docx
    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    tmp.write(file_bytes)
    tmp.flush()
    tmp.close()
    try:
        doc = Document(tmp.name)
        texts = []
        for p in doc.paragraphs:
            if p.text and p.text.strip():
                texts.append(p.text.strip())
        return "\n".join(texts)
    finally:
        os.unlink(tmp.name)

def insert_comment_into_docx_bytes(original_bytes: bytes, comment_locations: List[Dict[str, Any]]) -> bytes:
    """
    Add bracketed inline comments and a summary page to the end of the docx file.
    comment_locations: list of dicts:
      {
        "match_text": "...snippet matched...",
        "comment": "Text of the suggestion or rule citation",
        "severity": "High"/"Medium"/"Low",
        "context": "optional context"
      }
    Returns resulting docx as bytes.
    """
    tmp_in = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    tmp_in.write(original_bytes)
    tmp_in.flush()
    tmp_in.close()

    doc = Document(tmp_in.name)

    # Try to insert inline bracketed comments: for each paragraph that contains match_text, append comment
    for loc in comment_locations:
        match = loc.get("match_text", "").strip()
        comment = loc.get("comment", "")
        severity = loc.get("severity", "")
        inserted = False
        if not match:
            continue
        for p in doc.paragraphs:
            if match.lower() in p.text.lower():
                # append bracketed comment (simple)
                p.add_run(f"  [COMMENT ({severity}): {comment}]").italic = True
                inserted = True
        # if not found, add to end summary block instead
        if not inserted:
            # will be added later to summary page
            continue

    # Add a summary page with all comments and suggested actions
    doc.add_page_break()
    summary = doc.add_paragraph()
    summary.style = doc.styles['Normal']
    summary.add_run("Corporate Agent Review Summary\n").bold = True
    p = doc.add_paragraph()
    for i, loc in enumerate(comment_locations, start=1):
        doc.add_paragraph(f"{i}. Document snippet: {loc.get('match_text','(n/a)')}")
        doc.add_paragraph(f"   Issue: {loc.get('comment','')}")
        doc.add_paragraph(f"   Severity: {loc.get('severity','')}")
        doc.add_paragraph("")

    # Save to bytes
    tmp_out = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(tmp_out.name)
    tmp_out.flush()
    tmp_out.close()

    with open(tmp_out.name, "rb") as f:
        out_bytes = f.read()

    # cleanup
    os.unlink(tmp_in.name)
    os.unlink(tmp_out.name)
    return out_bytes

# ---------------- HELPERS: Doc type identification & checklist ----------------
def identify_doc_type_from_text(text: str) -> str:
    """Simple keyword-based doc type identification. Returns the best matching doc type or 'Unknown'."""
    text_low = text.lower()
    scores = {}
    for dtype, kws in DOC_KEYWORDS.items():
        for kw in kws:
            if kw in text_low:
                scores[dtype] = scores.get(dtype, 0) + 1
    if not scores:
        return "Unknown"
    # return highest scoring
    return max(scores.items(), key=lambda x: x[1])[0]

def determine_process_and_missing_docs(detected_types: List[str]) -> Tuple[str, List[str], List[str]]:
    """
    Simple approach:
      - If any of the checklist documents exist, assume "Company Incorporation".
      - Return (process_name, uploaded_doc_types, missing_docs)
    """
    proc = "Company Incorporation"
    required = CHECKLISTS.get(proc, [])
    uploaded = [t for t in detected_types if t != "Unknown"]
    missing = [d for d in required if d not in uploaded]
    return proc, uploaded, missing

# ---------------- HELPERS: Clause analysis using RAG + LLM ----------------
def analyze_snippet_with_llm(snippet: str, rag_contexts: List[Dict[str,Any]], llm: OllamaLLM) -> Dict[str, Any]:
    """
    Given a snippet of a document and RAG contexts (list of docs from the chroma index),
    call the LLM to determine if there's a red flag and to generate a suggested fix and citation.
    Returns a dict: { issue: str, severity: str, suggestion: str, citation: str }
    """
    # build prompt
    contexts_text = "\n\n---\n\n".join([c["document"] for c in rag_contexts])
    prompt = f"""
You are a legal-review assistant specialized in ADGM (Abu Dhabi Global Market) regulations.
Given the document snippet below, analyze whether it is compliant with ADGM corporate/company rules.
If there's a non-compliance or red-flag, respond in JSON with these keys:
- issue: short description of the problem.
- severity: High/Medium/Low.
- suggestion: a suggested replacement clause or corrective action (brief).
- citation: cite the ADGM law/regulation or section (if available) or the reference doc title.

Use the provided ADGM reference contexts (after "CONTEXT:") to ground your answer.

DOCUMENT SNIPPET:
\"\"\"{snippet}\"\"\"

CONTEXT:
\"\"\"{contexts_text}\"\"\"

Return only JSON.
"""
    # call the LLM (synchronous)
    try:
        # langchain-ollama's Ollama exposes the standard LangChain LLM interface .predict()
        result_text = llm.predict(prompt)
    except Exception as e:
        # fallback short default
        return {
            "issue": f"LLM call error: {e}",
            "severity": "Low",
            "suggestion": "",
            "citation": ""
        }

    # try parse JSON from result_text (best-effort)
    try:
        j = json.loads(result_text.strip())
        return j
    except Exception:
        # if LLM didn't return strict JSON, try to extract a JSON chunk
        import re
        m = re.search(r'\{.*\}', result_text, flags=re.S)
        if m:
            try:
                return json.loads(m.group(0))
            except Exception:
                pass
    # if all fails, return the LLM output as 'issue'
    return {
        "issue": result_text.strip()[:800],
        "severity": "Medium",
        "suggestion": "",
        "citation": ""
    }

# ---------------- MAIN PROCESSING: review uploaded docx files ----------------
def review_uploaded_docx_files(file_objs: List[Tuple[str, bytes]]) -> Dict[str, Any]:
    """
    file_objs is a list of tuples: (filename, bytes)
    Returns a dictionary with:
      - reviewed_files: list of tuples (filename, reviewed_bytes)
      - report: structured JSON-like dict summarizing missing docs and issues
    """
    # initialize LLM + embedding
    llm = OllamaLLM(model=LLM_MODEL)  # ensure you have this model available locally in Ollama
    embedder = OllamaEmbeddings(model=EMBED_MODEL)

    detected_types = []
    file_texts = {}
    results_for_report = []

    # 1) Extract text and identify types
    for fname, fbytes in file_objs:
        txt = extract_text_from_docx_bytes(fbytes)
        file_texts[fname] = txt
        dtype = identify_doc_type_from_text(txt)
        detected_types.append(dtype)

    # 2) Determine process & missing docs
    process, uploaded, missing = determine_process_and_missing_docs(detected_types)

    # 3) For each document, find suspicious phrases/clause heuristics
    reviewed_files = []
    for fname, fbytes in file_objs:
        txt = file_texts[fname]
        issues_found = []

        # Heuristic: split into paragraphs, analyze short paragraphs with LLM+RAG
        paragraphs = [p.strip() for p in txt.split("\n") if p.strip()]
        # limit paragraphs considered to top N to save time
        paragraph_sample = paragraphs[:200]  # cap to first 200 paragraphs

        for para in tqdm(paragraph_sample, desc=f"Analyzing {fname}", leave=False):
            # quick heuristics for obviously problematic phrases
            low = para.lower()
            heuristic_flag = False
            if "federal courts" in low or "uae federal" in low:
                heuristic_flag = True
            if "notwithstanding" in low and "adgm" not in low and "jurisdiction" in low:
                heuristic_flag = True
            if len(para) > 1000:  # long paragraph might hide issues
                heuristic_flag = True

            if heuristic_flag:
                # perform a RAG retrieval for context
                rag_ctxs = rag_retrieve(para, top_k=RAG_TOP_K)
                analysis = analyze_snippet_with_llm(para, rag_ctxs, llm)
                issue = {
                    "document": fname,
                    "section_snippet": para[:600],
                    "issue": analysis.get("issue", ""),
                    "severity": analysis.get("severity", "Medium"),
                    "suggestion": analysis.get("suggestion", ""),
                    "citation": analysis.get("citation", "")
                }
                issues_found.append(issue)

        # Build comment locations for insertion
        comment_locations = []
        for iss in issues_found:
            comment_locations.append({
                "match_text": iss["section_snippet"],
                "comment": f"{iss['issue']} Suggestion: {iss['suggestion']}. Citation: {iss['citation']}",
                "severity": iss["severity"]
            })

        # Insert comments and get reviewed bytes
        reviewed_bytes = insert_comment_into_docx_bytes(fbytes, comment_locations)
        reviewed_files.append((fname, reviewed_bytes))

        # add to report
        results_for_report.extend([
            {
                "document": iss["document"],
                "section": iss["section_snippet"][:200],
                "issue": iss["issue"],
                "severity": iss["severity"],
                "suggestion": iss["suggestion"],
                "citation": iss["citation"]
            } for iss in issues_found
        ])

    # final report
    report = {
        "process_detected": process,
        "documents_uploaded_count": len(file_objs),
        "required_documents_count": len(CHECKLISTS.get(process, [])),
        "uploaded_document_types": uploaded,
        "missing_documents": missing,
        "issues_found": results_for_report
    }

    return {
        "reviewed_files": reviewed_files,
        "report": report
    }

# ---------------- GRADIO UI ----------------
def gradio_handler(files):
    """
    files: list of temporary file dicts from gradio
    Returns downloadable zip of reviewed docs and JSON report
    """
    # files is a list of dicts with keys: name, data (or tuple)
    file_objs = []
    for f in files:
        # Gradio may give tempfile-like object; read as bytes
        if isinstance(f, tuple) or isinstance(f, list):
            # tuple (file_name, bytes)
            file_objs.append((f[0], f[1]))
        else:
            # gr.File returns a tempfile path in f.name
            with open(f.name, "rb") as fh:
                file_bytes = fh.read()
            file_objs.append((os.path.basename(f.name), file_bytes))

    result = review_uploaded_docx_files(file_objs)

    # save reviewed files to a temp folder and zip
    tmpd = tempfile.mkdtemp(prefix="reviewed_")
    reviewed_paths = []
    for fname, b in result["reviewed_files"]:
        outp = os.path.join(tmpd, f"reviewed_{fname}")
        with open(outp, "wb") as fh:
            fh.write(b)
        reviewed_paths.append(outp)

    # write JSON report
    report_path = os.path.join(tmpd, "report.json")
    with open(report_path, "w", encoding="utf-8") as f:
        json.dump(result["report"], f, ensure_ascii=False, indent=2)

    # build a zip to return
    import zipfile
    zip_path = os.path.join(tmpd, "corporate_agent_review.zip")
    with zipfile.ZipFile(zip_path, "w") as z:
        for p in reviewed_paths:
            z.write(p, arcname=os.path.basename(p))
        z.write(report_path, arcname="report.json")

    # return the zip and the JSON for display
    return zip_path, result["report"]


def main_ui():
    title = "ADGM Corporate Agent — Document Intelligence (Gradio demo)"
    description = "Upload .docx documents for review. The agent will check checklist compliance, insert inline comments, and return a reviewed .docx + JSON report."

    with gr.Blocks() as demo:
        gr.Markdown(f"## {title}\n\n{description}")
        file_in = gr.File(label="Upload .docx files", file_count="multiple", file_types=[".docx"])
        run_btn = gr.Button("Review Documents")
        output_zip = gr.File(label="Download reviewed files (zip)")
        output_json = gr.JSON(label="JSON report")
        run_btn.click(fn=gradio_handler, inputs=[file_in], outputs=[output_zip, output_json])

    demo.launch(share = True)

# ---------------- RUN ----------------
if __name__ == "__main__":
    main_ui()
